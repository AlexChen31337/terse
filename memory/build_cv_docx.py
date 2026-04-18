"""Build Bowen Li's polished CV as a professional Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x33, 0x5C)   # deep navy – headings
TEAL   = RGBColor(0x00, 0x7A, 0x87)   # teal accent – name / section rules
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black body
MID    = RGBColor(0x44, 0x44, 0x55)   # mid-grey sub-text
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# Tuple versions for XML attribute use
TEAL_T = (0x00, 0x7A, 0x87)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, italic=False, color=DARK):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color

def para_spacing(p, before=0, after=4):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def add_hr(doc, color=None):
    """Insert a thin coloured horizontal rule."""
    if color is None:
        color = TEAL_T
    p = doc.add_paragraph()
    para_spacing(p, before=2, after=6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_name(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=0, after=2)
    run = p.add_run("BOWEN LI")
    set_font(run, 28, bold=True, color=TEAL)

def add_tagline(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=0, after=2)
    run = p.add_run("AI System Architect  |  Principal Engineer  |  Agentic AI Specialist")
    set_font(run, 11, italic=True, color=NAVY)

def add_contact(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=0, after=8)
    run = p.add_run(
        "Kellyville, NSW 2155  ·  0430 830 888  ·  bowen.li@outlook.com.au\n"
        "Australian Citizen  ·  Baseline Security Clearance"
    )
    set_font(run, 9.5, color=MID)

def add_section_header(doc, title):
    add_hr(doc)
    p = doc.add_paragraph()
    para_spacing(p, before=4, after=3)
    run = p.add_run(title.upper())
    set_font(run, 11, bold=True, color=NAVY)

def add_body(doc, text, size=10, color=DARK, bold=False, italic=False, before=0, after=3):
    p = doc.add_paragraph()
    para_spacing(p, before=before, after=after)
    run = p.add_run(text)
    set_font(run, size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    para_spacing(p, before=0, after=2)
    run = p.add_run(text)
    set_font(run, size, color=DARK)

def add_role(doc, title, org, location, period):
    p = doc.add_paragraph()
    para_spacing(p, before=6, after=1)
    r1 = p.add_run(f"{title}  ")
    set_font(r1, 10.5, bold=True, color=NAVY)
    r2 = p.add_run(f"— {org}")
    set_font(r2, 10.5, bold=True, color=TEAL)
    p2 = doc.add_paragraph()
    para_spacing(p2, before=0, after=3)
    r3 = p2.add_run(f"{location}  |  {period}")
    set_font(r3, 9, italic=True, color=MID)

def add_subheading(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=5, after=2)
    run = p.add_run(text)
    set_font(run, 10, bold=True, color=NAVY)

# ══════════════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

add_name(doc)
add_tagline(doc)
add_contact(doc)

# ── SUMMARY ──────────────────────────────────────────────────────────────────
add_section_header(doc, "Profile")
add_body(doc,
    "Engineering leader with 14+ years designing and scaling enterprise-grade platforms — "
    "and 3+ years at the frontier of agentic AI systems, building autonomous agents, "
    "multi-agent orchestration frameworks, and production-grade LLM pipelines. One of a rare "
    "breed who combines deep systems engineering with modern AI-native development: "
    "context-driven architecture, tool-use harness engineering, and self-evolving agent frameworks.",
    size=10, after=3)
add_body(doc,
    "Delivered enterprise outcomes for Coles, ASX, CBA, Accenture (Federal Government), "
    "Woolworths, Samsung, Optus, Telstra, and large-scale fintech clients. Now applying that "
    "foundation to the next generation of intelligent, autonomous systems.",
    size=10, after=3)

# ── CORE COMPETENCIES ────────────────────────────────────────────────────────
add_section_header(doc, "Core Competencies")

add_subheading(doc, "🤖  Agentic AI & LLM Systems")
bullets_agent = [
    "Multi-agent orchestration: Planner/Builder/Reviewer pipelines, agent-to-agent communication, task routing, session affinity",
    "Context-driven architecture: long-context retrieval, memory tiering (working/episodic/long-term), WAL-based state persistence, compaction-safe agent design",
    "Tool-use & harness engineering: agent tool loops, tool registries, capability scaffolding, MCP-compatible tool surfaces",
    "LLM fine-tuning: LoRA, QLoRA, supervised fine-tuning, RLHF, DPO on open-source models",
    "RAG pipelines: LangChain, LangGraph, hybrid vector+keyword retrieval, reranking, knowledge graph integration",
    "Agent runtimes: custom Go-based runtimes, EvoClaw (self-evolving agent framework), OpenClaw, CrewAI patterns",
    "Prompt engineering & evals: chain-of-thought, structured outputs, automated regression suites, drift detection",
    "MLOps: GPU-backed inference (CUDA, multi-GPU), model versioning, A/B testing, deployment pipelines",
    "Tooling: Ollama, HuggingFace, OpenAI API, Anthropic API, NVIDIA NIM, ComfyUI, Dify",
]
for b in bullets_agent:
    add_bullet(doc, b)

add_subheading(doc, "🏗️  Systems Architecture")
for b in [
    "Microservices, event-driven systems (Kafka, MQTT, NATS), distributed state machines",
    "Enterprise API platforms: REST, GraphQL, gRPC, WebSocket at scale",
    "Blockchain / L1 design: Substrate-based agent-native chains, NPoS consensus, on-chain agent identity (DIDs), reputation pallets",
    "Infrastructure as Code: Terraform, Pulumi",
    "Systems design for fintech, trading automation, and high-security government workloads",
]:
    add_bullet(doc, b)

add_subheading(doc, "💻  Modern Software Engineering")
for b in [
    "Context-driven development: spec-first, documentation-driven, AI-assisted coding workflows",
    "Languages: Go, Python, TypeScript/Node.js, Rust, Java Spring Boot, .NET/C#",
    "Frontend: React, Next.js, Angular, Vue.js, TypeScript",
    "AI-native tooling: GitHub Copilot, Claude Code, cursor-driven refactoring, LLM-in-the-loop CI/CD",
    "Testing: TDD/BDD, property-based testing, coverage-gated CI (90%+ standard), fuzz testing",
]:
    add_bullet(doc, b)

add_subheading(doc, "☁️  Cloud & DevOps")
for b in [
    "AWS: ECS, EKS, Lambda, API Gateway, DynamoDB, S3, RDS, Bedrock",
    "Azure: App Services, AKS, CosmosDB, Azure OpenAI Service",
    "GCP: GKE, Cloud Run, BigQuery, Vertex AI",
    "Containers & Orchestration: Docker, Kubernetes, Helm",
    "CI/CD: GitHub Actions, Jenkins, Azure DevOps — including AI model deployment pipelines",
    "Observability: Datadog, Prometheus, Grafana, Sentry, OpenTelemetry",
]:
    add_bullet(doc, b)

# ── EXPERIENCE ───────────────────────────────────────────────────────────────
add_section_header(doc, "Professional Experience")

add_role(doc, "Senior AI Forward Deployed Engineer", "NOES", "Sydney, NSW", "Jan 2026 – Present")
add_body(doc, "Leading AI system architecture and field engineering for enterprise clients transitioning to agentic AI workflows.", size=10, after=2)
for b in [
    "Designing multi-agent systems with tool-use harnesses, long-context memory, and autonomous task execution",
    "Building enterprise RAG pipelines with LangGraph, knowledge graphs, and hybrid retrieval over proprietary corpora",
    "Advising on LLM selection, fine-tuning strategy, and production deployment for regulated environments",
    "Establishing MLOps foundations: GPU inference clusters, model versioning, evaluation harnesses, drift monitoring",
]:
    add_bullet(doc, b)

add_role(doc, "Senior Full Stack Developer & AI Engineer", "Future Secure AI", "Sydney, NSW", "Nov 2024 – Dec 2025")
for b in [
    "Architected LLM-powered enterprise applications with RAG, knowledge graphs, and long-context retrieval across high-security environments",
    "Built agentic pipelines with tool-use loops, structured output parsing, and automated fallback handling",
    "Designed scalable API layers (GraphQL/REST) serving agent orchestration backends",
    "Implemented full CI/CD for model + software co-deployment (GitHub Actions + Docker + K8s)",
    "Set up comprehensive observability: Datadog, Prometheus, Grafana, OpenTelemetry",
]:
    add_bullet(doc, b)

add_role(doc, "Principal Engineer / Full Stack / AI Engineer", "Uno Digit", "Sydney, NSW", "Jan 2021 – Present")

add_subheading(doc, "AI & Agent Engineering")
for b in [
    "Designed and shipped EvoClaw — a self-evolving agent framework in Go featuring skill distillation, recursive self-improvement loops, and multi-platform deployment (Android, iOS, WASM, ClawHub)",
    "Built ClawChain — a Substrate-based L1 blockchain for autonomous agents: NPoS consensus, on-chain DIDs, reputation pallets, service markets, and IBC-lite cross-chain messaging",
    "Implemented agentic CI/CD pipelines: autonomous PR review agents, CI health monitors, and code quality agents operating across multiple production repos",
    "Built domain-specific LLM solutions with custom RAG pipelines, tool registries, and tiered agent memory systems",
]:
    add_bullet(doc, b)

add_subheading(doc, "Backend & Systems")
for b in [
    "Multi-tenant backend platforms: Node.js, Go, Java, C# — API-first, event-driven architecture",
    "High-availability systems for Department of Agriculture, Coles, ASX, MintPayments (Fintech)",
    "Fintech-grade auth: OAuth2, JWT, encryption, compliance-aware audit logging",
]:
    add_bullet(doc, b)

add_subheading(doc, "Highlighted Projects")
for b in [
    "Coles — Checkout and customer feedback platform (high-traffic, PCI-compliant)",
    "ASX — Secure online submissions platform with audited third-party integrations",
    "MintPayments — Fintech CRM/portal: React + Node.js + AWS",
]:
    add_bullet(doc, b)

add_role(doc, "Frontend Development Associate Manager", "Accenture", "Sydney, NSW", "Jun 2017 – Jan 2021")
for b in [
    "Australian Federal Government — My Health Record: full accessibility remediation for a platform serving 23M+ Australians",
    "Woolworths — Demand forecasting platform: React + Django + GCP",
    "ANZ Bank — Internal conversational AI chatbot: Dialogflow, React, Node.js",
    "Led large-scale digital transformations using React, Angular, Vue, Node.js, AWS, Azure",
    "Production support for high-traffic systems; CI/CD via Jenkins, GitHub Actions, Azure Pipelines",
]:
    add_bullet(doc, b)

add_role(doc, "Principal Engineer / Tech Lead", "Hinterlands", "Sydney, NSW", "Aug 2012 – Jun 2017")
for b in [
    "Designed REST/GraphQL API platforms and scalable backends handling millions of monthly interactions",
    "Clients: Optus (Yes Crowd), CBA (Online Community), Samsung UK (E-commerce), Telstra (Community Platform), Bunnings, Vodafone, MYOB, HR Block",
    "Deployed on AWS with Infrastructure-as-Code and early CI/CD adoption",
]:
    add_bullet(doc, b)

# ── EDUCATION ────────────────────────────────────────────────────────────────
add_section_header(doc, "Education")
add_body(doc, "Master of Science (Engineering)  —  Politecnico di Torino, Italy", size=10, bold=True, after=2)
add_body(doc, "Bachelor of Science (Engineering)  —  Wuhan University of Technology, China", size=10, bold=True, after=3)

# ── CERTIFICATIONS ────────────────────────────────────────────────────────────
add_section_header(doc, "Certifications")
for b in [
    "AWS Solutions Architect – Associate",
    "AWS AI Practitioner",
    "Zend PHP Certified Engineer",
]:
    add_bullet(doc, b)

# ── ADDITIONAL ────────────────────────────────────────────────────────────────
add_section_header(doc, "Publications & Community")
for b in [
    "React Educator — Taught advanced React courses at RMIT Online",
    "Author — Published two books on React.js",
    "Open-Source Contributor — React core and multiple frontend libraries",
    "Technical Interviewer — Senior screening for React and full-stack engineering roles",
]:
    add_bullet(doc, b)

# ── DIFFERENTIATOR ────────────────────────────────────────────────────────────
add_section_header(doc, "What Sets Me Apart")
add_body(doc,
    "Most engineers can build with AI. Few can build AI systems that build themselves.",
    size=10, italic=True, bold=True, color=TEAL, before=2, after=4)
add_body(doc,
    "I operate at the intersection of deep systems engineering and frontier AI architecture — "
    "writing Go runtimes for autonomous agents, designing on-chain reputation systems for AI "
    "identities, and shipping multi-agent pipelines that run autonomously in production. "
    "I don't just integrate LLM APIs: I architect the harnesses, memory systems, tool registries, "
    "and self-improvement loops that make agents reliable, cost-efficient, and genuinely autonomous.\n\n"
    "If you're building the next generation of AI-native products, I'm the engineer "
    "who's already been living there.",
    size=10, after=3)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = "/home/bowen/.openclaw/workspace/memory/Bowen_Li_CV.docx"
doc.save(out)
print(f"Saved: {out}")
